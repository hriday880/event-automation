import streamlit as st
import tempfile
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_report_docx(event_details, photo_files):
    """Generates the Word document report based on the provided details and photos."""
    doc = Document()
    
    # Add Logo to the Header so it appears on every page (like the empty page format)
    try:
        section = doc.sections[0]
        header = section.header
        # The header already has an empty paragraph by default
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.add_run()
        header_run.add_picture("logo.jpg", width=Inches(1.8))
    except Exception as e:
        st.warning(f"Could not load logo into header: {e}")
        
    # Title
    title = doc.add_heading(f"{event_details['Event Name']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Details Table-like layout using tabs
    details = [
        ("Academic Year", event_details['Academic Year']),
        ("Name of the Activity", event_details['Event Name']),
        ("Date of the Activity", event_details['Date'].strftime("%d.%m.%Y") if event_details['Date'] else ""),
        ("Venue", event_details['Venue']),
        ("Organized By", event_details['Organized By']),
        ("No. of Participants", str(event_details['No. of Participants'])),
        ("Brief Report", event_details['Brief Report']),
    ]
    
    for label, value in details:
        p = doc.add_paragraph()
        p.add_run(f"{label}:\t").bold = True
        p.add_run(str(value))

    doc.add_page_break()
    
    # Add Photos
    p = doc.add_paragraph()
    p.add_run("Photograph of the Event (1 geotagged compulsory)").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for photo in photo_files:
        # Save uploaded file to a temp file so python-docx can read it
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
            tmp_file.write(photo.getvalue())
            tmp_path = tmp_file.name
        
        try:
            doc.add_picture(tmp_path, width=Inches(6.0))
        except Exception as e:
            st.warning(f"Could not insert a photo into docx: {e}")
        finally:
            os.remove(tmp_path)
            
    # Save the generated docx to a BytesIO object for downloading
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- Streamlit UI ---
st.set_page_config(page_title="Event Report Generator", layout="centered")
st.title("📝 Event Report Generator")
st.markdown("Fill in the details below to generate your formatted Event Report Document. You can then download it and upload it to Google Drive manually.")

with st.form("event_form"):
    st.subheader("1. Event Details")
    col1, col2 = st.columns(2)
    with col1:
        academic_year = st.text_input("Academic Year", value="2026-27")
        event_name = st.text_input("Event Name", placeholder="e.g. FUSIP Ice Breakers")
        event_date = st.date_input("Date of Activity")
        event_type = st.selectbox("Event Type", ["Cultural", "Club", "Academic", "Sports", "Other"])
    with col2:
        venue = st.text_input("Venue", placeholder="e.g. Kalidas")
        organized_by = st.text_input("Organized By", value="DotSlash")
        no_of_participants = st.number_input("No. of Participants", min_value=1, value=50)
        
    brief_report = st.text_area("Brief Report", height=200, placeholder="Write the report content here...")
    
    st.subheader("2. Upload Photos")
    st.markdown("📸 **Photos** (at least 3-4, landscape, geotagged)")
    photos = st.file_uploader("Upload Event Photos", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
    
    submit = st.form_submit_button("Generate Report Document 📄")

if submit:
    if not event_name or not brief_report:
        st.error("Please fill in the Event Name and Brief Report.")
    elif not photos or len(photos) < 3:
        st.error("Please upload at least 3 photos.")
    else:
        with st.spinner("Generating Document..."):
            event_details = {
                "Academic Year": academic_year,
                "Event Name": event_name,
                "Date": event_date,
                "Venue": venue,
                "Organized By": organized_by,
                "No. of Participants": no_of_participants,
                "Brief Report": brief_report
            }
            
            # Generate the docx in memory
            doc_file = generate_report_docx(event_details, photos)
            
            st.success("✅ Document Generated Successfully!")
            st.download_button(
                label="⬇️ Download Event Report (.docx)",
                data=doc_file,
                file_name=f"{event_name}_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            st.info("""
            **Next Steps for Secretary:**
            1. Download the document above.
            2. Upload it to the `1. Club Activity Reports` folder in Google Drive.
            3. Create a folder in `2. Picture Repository` and upload the raw photos there.
            4. Duplicate the Participant Template in `3. Organiser & Participant Repository` and fill it.
            5. Update the `Club & Cultural Activities 2026-27` Spreadsheet with the links.
            """)
